from docx import Document
from docx.shared import Cm

doc = Document('test.docx')
print(doc.paragraphs)  # 段落

for paragraph in doc.paragraphs:
    print(paragraph.text)  # 列出内容

# 提取第二个段落文字块
print('*************提取文字块***********************')
paragraph = doc.paragraphs[1]
runs = paragraph.runs
for run in runs:
    print(run.text)

print('添加标题')
doc.add_heading('一级标题', level=1)

print('添加段落')
paragraph1 = doc.add_paragraph('这是一个段落')
paragraph2 = doc.add_paragraph('这又是一个段落')

print('添加文字块')
paragraph3 = doc.add_paragraph()
paragraph3.add_run('加粗').bold = True
paragraph3.add_run('普通')
paragraph3.add_run('斜体').italic = True

print('添加分页')
doc.add_page_break()

print('添加图片')
doc.add_picture('test.png')

print('添加图片，给定高度或宽度')
doc.add_picture('test.png', width=Cm(5), height=Cm(5))

print('添加表格')
records = [
    ['学号', '姓名', '成绩'],
    [101, '李雷', 95],
    [102, '韩梅梅', 100],
    [103, '马冬梅', 98]
]

table = doc.add_table(rows=4, cols=3)
for row in range(4):
    cells = table.rows[row].cells
    for col in range(3):
        cells[col].text = str(records[row][col])

print('保存')
doc.save('test.docx')


# 对文字字体样式进行修改
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx import Document

doc = Document('test.docx')

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.bold = True
        run.font.italic = True
        run.font.underline = True
        run.font.strike = True  # 删除线
        run.font.shadow = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(255, 255, 0)

        run.font.name = '微软雅黑'
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), '微软雅黑')

doc.save('test1.docx')

# 段落样式的修改
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 行间距
paragraph.paragraph_format.line_spacing = 2.0

# 段前与段后间距
paragraph.paragraph_format.space_before = Pt(12)
paragraph.paragraph_format.space_after = Pt(12)


